import streamlit as st
import PyPDF2
from google import genai
from docx import Document
from io import BytesIO
import time

# --- 1. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Robotmaster OLP Auditor", page_icon="🤖", layout="wide")

# CSS "Nuclear Black" - Forçando Preto Absoluto contra o tema do Streamlit
st.markdown("""
    <style>
    .report-container {
        background-color: #ffffff;
        padding: 45px;
        border-radius: 12px;
        border: 1px solid #e0e0e0;
        box-shadow: 0px 8px 25px rgba(0,0,0,0.08);
        line-height: 1.7;
        font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    }
    /* Regras agressivas para derrubar o azul do Streamlit */
    div[data-testid="stMarkdownContainer"] p, 
    div[data-testid="stMarkdownContainer"] li, 
    div[data-testid="stMarkdownContainer"] span {
        color: #000000 !important; 
    }
    div[data-testid="stMarkdownContainer"] h1, 
    div[data-testid="stMarkdownContainer"] h2, 
    div[data-testid="stMarkdownContainer"] h3,
    div[data-testid="stMarkdownContainer"] h4 {
        color: #000000 !important; font-weight: 900 !important;
    }
    .report-container h2 {
        text-transform: uppercase; border-bottom: 3px solid #000000; padding-bottom: 8px; font-size: 22px; letter-spacing: 1px; margin-top: 30px;
    }
    .report-header { text-align: center; margin-bottom: 30px; padding-bottom: 20px; border-bottom: 1px solid #000000; }
    .stButton>button { border-radius: 6px !important; }
    </style>
    """, unsafe_allow_html=True)

# Inicializar Session State (A Memória do App)
if "history" not in st.session_state: st.session_state.history = []
if "current_report" not in st.session_state: st.session_state.current_report = None
if "chat_session" not in st.session_state: st.session_state.chat_session = None
if "mensagens_chat" not in st.session_state: st.session_state.mensagens_chat = []
# SOLUÇÃO DO ERRO DO CHAT: Guardar a conexão na memória!
if "genai_client" not in st.session_state: st.session_state.genai_client = None

# --- 2. SIDEBAR ---
with st.sidebar:
    try: st.image("LOGO_Robotmaster_RGB.png", use_container_width=True)
    except: st.markdown("<h2 style='color:#000000; font-weight:900; text-align:center;'>ROBOTMASTER</h2>", unsafe_allow_html=True)
    st.markdown("---")
    st.title("🕒 Recent Audits")
    for idx, hist in enumerate(reversed(st.session_state.history)):
        if st.button(f"📄 Audit {hist['time']}", key=f"hist_{idx}", use_container_width=True):
            st.session_state.current_report = hist['content']

# --- 3. INTERFACE PRINCIPAL ---
st.markdown("<h3 style='text-align: center; color: #000000; font-weight: 900;'>ADVANCED OLP ENGINEERING AUDITOR</h3>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #555;'>Upload the Client RFP for deep kinematic and ROI analysis.</p>", unsafe_allow_html=True)

uploaded_file = st.file_uploader("", type="pdf")

if st.button("🚀 EXECUTE ROBOTMASTER V7 ANALYSIS", type="primary", use_container_width=True):
    if uploaded_file:
        with st.status("⚙️ Initializing Anbu Intelligence Engine...", expanded=True) as status:
            try:
                st.write("📄 Extracting technical data from RFP...")
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                pdf_text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
                
                st.write("🤖 Connecting to Gemini Neural Network...")
                # Cria a conexão e GUARDA na memória para o chat não cair depois
                st.session_state.genai_client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
                
                st.write("📐 Mapping Kinematics & V7 Architecture...")
                
                # PROMPT CORRIGIDO PARA O ROI FICAR GIGANTE
                system_instruction = """
                You are an Elite Robotics Solutions Architect & Senior Robotmaster V7 OLP Specialist.
                Your mission is to analyze client RFPs and deliver a highly technical, persuasive, and data-driven Engineering & Integration Report.

                FORMAT STRICTLY USING THE FOLLOWING SECTIONS (Use '##' for titles):

                ## 1. EXECUTIVE SUMMARY & PROCESS OVERVIEW
                Provide a high-level summary of the client's manufacturing process. Identify the core robotic application and the primary bottleneck.

                ## 2. MACHINERY & KINEMATICS ANALYSIS
                Analyze the robots, positioners, rails, and tooling mentioned. Highlight any kinematic complexities like external axes management.

                ## 3. ROBOTMASTER V7 SOLUTION ARCHITECTURE
                Map the exact Robotmaster V7 modules required. Explain WHY they are needed.

                ## 4. ADVANCED R.O.I. & METRICS COMPARISON
                Provide a COMPREHENSIVE and EXTENSIVE list of ALL relevant ROI metrics (minimum 8 to 12 items). Analyze every possible efficiency gain.
                Use exactly this format for EACH item (NO TABLES):
                - **[Metric Name]:** [Manual/Teach Pendant] vs [Robotmaster V7] ➔ [Quantifiable Savings/Impact]

                ## 5. RISK ASSESSMENT & MITIGATION
                Identify technical risks and explain how Robotmaster's Optimization tools automatically mitigate them.

                ## 6. EXECUTIVE RECOMMENDATION
                A strong, closing paragraph convincing the client that Robotmaster is the absolute best software investment.

                TONE: Highly technical, authoritative, yet persuasive. Use precise industrial robotics terminology. DO NOT USE TABLES.
                """
                
                response = st.session_state.genai_client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=f"{system_instruction}\n\nDocument:\n{pdf_text}"
                )
                
                st.session_state.current_report = response.text
                st.session_state.history.append({"time": time.strftime("%H:%M"), "content": response.text})
                
                st.write("✅ Initializing Copilot Chat...")
                instrucao_chat = f"Atue como meu assistente de pré-vendas Robotmaster. Baseie-se nesta RFP do cliente:\n\n{pdf_text}"
                
                # O chat agora usa a conexão guardada na memória!
                st.session_state.chat_session = st.session_state.genai_client.chats.create(
                    model='gemini-2.5-flash',
                    config={"system_instruction": instrucao_chat}
                )
                
                st.session_state.mensagens_chat = [
                    {"role": "assistant", "content": "Olá! O relatório foi gerado na aba ao lado. Você pode me pedir para resumir partes, criar um e-mail para o cliente ou tirar dúvidas sobre o documento original!"}
                ]
                
                status.update(label="Analysis Successfully Completed!", state="complete", expanded=False)
                
            except Exception as e: 
                status.update(label="Critical Error Encountered", state="error", expanded=False)
                st.error(f"Error details: {e}")
    else:
        st.warning("⚠️ Please upload a PDF file first.")

# Função para exportar o chat em Word (.docx)
def gerar_word_chat(mensagens):
    doc = Document()
    doc.add_heading('Histórico do Copiloto - Robotmaster', 0)
    for msg in mensagens:
        autor = "Sales Engineer" if msg["role"] == "user" else "Anbu AI"
        doc.add_heading(autor, level=2)
        doc.add_paragraph(msg['content'])
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- 4. EXIBIÇÃO EM ABAS (RELATÓRIO E CHAT) ---
if st.session_state.current_report:
    st.success("✅ **Auditoria concluída com sucesso!** Navegue pelas DUAS ABAS abaixo.")
    
    aba_relatorio, aba_chat = st.tabs(["📑 1. ENGINEERING REPORT", "💬 2. AI COPILOT CHAT"])
    
    # === ABA 1: RELATÓRIO ===
    with aba_relatorio:
        st.markdown('<div class="report-container">', unsafe_allow_html=True)
        st.markdown('<div class="report-header"><h4>CONFIDENTIAL ENGINEERING REPORT</h4><p style="color:#000000; font-weight:bold;">GENERATED BY ROBOTMASTER AI</p></div>', unsafe_allow_html=True)
        st.markdown(st.session_state.current_report)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        u_id = str(int(time.time()))
        c1, c2 = st.columns(2)
        with c1: 
            doc = Document()
            doc.add_heading('Robotmaster Engineering Report', 0)
            doc.add_paragraph(st.session_state.current_report)
            target = BytesIO()
            doc.save(target)
            st.download_button("📝 Download Editable Word Doc (.docx)", target.getvalue(), f"Robotmaster_Audit_{u_id}.docx", key=f"w_{u_id}", use_container_width=True, type="primary")
        with c2: 
            st.link_button("📧 Send to Sales Team", "mailto:sales@robotmaster.com", use_container_width=True)

    # === ABA 2: CHAT ===
    with aba_chat:
        st.caption("Converse com a Anbu AI sobre a RFP ou peça para redigir e-mails comerciais.")
        
        # O botão de exportar chat agora gera um DOCX!
        if st.session_state.mensagens_chat:
            word_chat_data = gerar_word_chat(st.session_state.mensagens_chat)
            st.download_button(
                label="📥 Export Chat History (.docx)",
                data=word_chat_data,
                file_name="Robotmaster_Chat_History.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        st.divider()
        
        for msg in st.session_state.mensagens_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if prompt := st.chat_input("Ex: Escreva um e-mail curto destacando a solução de soldagem..."):
            st.session_state.mensagens_chat.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
                
            with st.chat_message("assistant"):
                with st.spinner("Pensando..."):
                    try:
                        resposta_chat = st.session_state.chat_session.send_message(prompt)
                        texto_resposta = resposta_chat.text
                    except Exception as e:
                        texto_resposta = f"⚠️ Erro de conexão com a IA: {e}. Por favor, clique no botão principal de executar análise novamente."
                    st.markdown(texto_resposta)
            st.session_state.mensagens_chat.append({"role": "assistant", "content": texto_resposta})